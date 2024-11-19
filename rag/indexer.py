import os
import logging
from pathlib import Path
from sqlalchemy import create_engine, text
from llama_index.vector_stores.postgres import PGVectorStore
from llama_index.embeddings.openai import OpenAIEmbedding
from llama_index.core.storage.docstore import SimpleDocumentStore
from llama_index.core.ingestion import IngestionPipeline
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core import Document
from sqlalchemy import make_url


logger = logging.getLogger("indexer")
logging.basicConfig(level=logging.INFO)

# Supported file types and their readers
file_extensions = [".pdf", ".docx", ".txt", ".md", ".html"]
embedding_model = "text-embedding-ada-002"
def create_index_table(db_url, index_table_name):
    """Create the index table if it does not exist."""
    logger.info(f"Attempting to create table: data_rag_{index_table_name}")
    engine = create_engine(db_url)
    create_table_query = f"""
    CREATE EXTENSION IF NOT EXISTS vector;

    CREATE TABLE IF NOT EXISTS data_rag_{index_table_name} (
        id UUID DEFAULT gen_random_uuid() PRIMARY KEY,
        metadata_ JSONB NOT NULL,
        node_id TEXT NOT NULL,
        embedding VECTOR(1536)
    );
    """
    try:
        with engine.connect() as connection:
            connection.execute(text(create_table_query))
            logger.info(f"Table data_rag_{index_table_name} created or already exists.")
    except Exception as e:
        logger.error(f"Error creating table: {e}")



def make_db_url(db):
    """Construct the database URL from the configuration dictionary."""
    return f"postgresql://{db['username']}:{db['password']}@{db['hostname']}:{db['port']}/{db['dbname']}"

def fetch_existing_content(db_url, index_table_name):
    """Fetch existing content metadata from the index table."""
    try:
        logger.info(f"Fetching existing content from table: data_rag_{index_table_name}")
        engine = create_engine(db_url)
        with engine.connect() as connection:
            query = text(f"SELECT id, metadata_, node_id FROM data_rag_{index_table_name};")
            result = connection.execute(query)
            rows = result.fetchall()
            columns = result.keys()
            content = [dict(zip(columns, row)) for row in rows]
        return content
    except Exception as e:
        logger.error(f"Error fetching content: {e}")
        return []

def delete_existing_content(db_url, index_table_name, file_names):
    """Delete outdated content based on file names."""
    try:
        engine = create_engine(db_url)
        with engine.begin() as connection:
            delete_query = text(f"""
                                    DELETE FROM data_rag_{index_table_name}
                                    WHERE metadata_->>'file_name' IN :file_names;
                                """)
            connection.execute(delete_query, {'file_names': tuple(file_names)})
            logger.info(f"Deleted records for files: {file_names}")
    except Exception as e:
        logger.error(f"Error deleting content: {e}")

def load_files_from_folder(folder_path, chunk_size, chunk_overlap):
    """Load documents from the specified folder."""
    documents = []
    for root, _, files in os.walk(folder_path):
        for file in files:
            ext = Path(file).suffix.lower()
            if ext not in file_extensions:
                continue
            file_path = os.path.join(root, file)
            metadata = {
                        'file_name': file,
                        'file_size': os.path.getsize(file_path),
                        'last_modified_date': os.path.getmtime(file_path),
                        'chunk_size': chunk_size,
                        'chunk_overlap': chunk_overlap
                    }

            try:
                if ext == ".txt":
                    with open(file_path, 'r', encoding='utf-8') as f:
                        text = f.read()
                    documents.append(Document(text=text, extra_info=metadata))
                elif ext == ".pdf":
                    import PyPDF2
                    text = ""
                    with open(file_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        for page in pdf_reader.pages:
                            text += page.extract_text() or ""
                    documents.append(Document(text=text, extra_info=metadata))
                elif ext == ".docx":
                    from docx import Document as DocxDocument
                    doc = DocxDocument(file_path)
                    text = "\n".join([para.text for para in doc.paragraphs])
                    documents.append(Document(text=text, extra_info=metadata))
                elif ext in [".html", ".htm"]:
                    from bs4 import BeautifulSoup
                    with open(file_path, 'r', encoding='utf-8') as f:
                        soup = BeautifulSoup(f, 'html.parser')
                    text = soup.get_text(separator="\n")
                    documents.append(Document(text=text, extra_info=metadata))
                logger.info(f"Loaded file: {file}")
            except Exception as e:
                logger.error(f"Error loading file {file}: {e}")
    return documents

def compare_metadata(existing_content, incoming_metadata):
    """Compare metadata to determine which files need reindexing."""
    existing_files = {item['metadata_']['file_name']: item for item in existing_content}
    incoming_files = {item['file_name']: item for item in incoming_metadata}

    updated_files = []
    for file_name, incoming_meta in incoming_files.items():
        existing_meta = existing_files.get(file_name, {}).get('metadata_', {})
        # Check for changes in file size, last modified date, chunk size, or chunk overlap
        if (existing_meta.get('file_size') != incoming_meta['file_size'] or
            existing_meta.get('last_modified_date') != incoming_meta['last_modified_date'] or
            existing_meta.get('chunk_size') != incoming_meta['chunk_size'] or
            existing_meta.get('chunk_overlap') != incoming_meta['chunk_overlap']):
            updated_files.append(file_name)
    return updated_files


def run_indexer(local_folder_path, index_table_name, chunk_size, chunk_overlap):
    """Main function to run the indexer."""
    db_url = make_url(os.getenv("VECTOR_DATABASE_URL"))
    try:
        logger.info("Connecting to vector store with URL parameters:")
        logger.info(f"Database: {db_url.database}, Host: {db_url.host}, User: {db_url.username}, Port: {db_url.port}")
        
        vector_store = PGVectorStore.from_params(
            database=db_url.database,
            host=db_url.host,
            password=db_url.password,
            port=db_url.port,
            user=db_url.username,
            table_name=f"rag_{index_table_name}",
            embed_dim=1536,
        )
        logger.info("Vector store initialized successfully.")
    except Exception as e:
        logging.info(f"Error initializing vector store for index '{index_table_name}': {e}")
        return None
    create_index_table(db_url, index_table_name)

    # Fetch existing content
    existing_content = fetch_existing_content(db_url, index_table_name)

    # Load new documents from the local folder
    documents = load_files_from_folder(local_folder_path, chunk_size, chunk_overlap)
    if not documents:
        logger.info("No documents found to index.")
        return

    # Extract metadata and compare
    incoming_metadata = [
            {
                'file_name': doc.extra_info['file_name'],
                'file_size': doc.extra_info['file_size'],
                'last_modified_date': doc.extra_info['last_modified_date'],
                'chunk_size': doc.extra_info.get('chunk_size'),
                'chunk_overlap': doc.extra_info.get('chunk_overlap')
            }
            for doc in documents
        ]

    files_to_reindex = compare_metadata(existing_content, incoming_metadata)

    if not files_to_reindex:
        logger.info("No changes detected in the files. Skipping indexing.")
        return

    # Delete outdated content
    delete_existing_content(db_url, index_table_name, files_to_reindex)

    # Filter documents for reindexing
    documents_to_reindex = [doc for doc in documents if doc.extra_info['file_name'] in files_to_reindex]

    # Run the ingestion pipeline
    pipeline = IngestionPipeline(
        transformations=[
            SentenceSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap),
            OpenAIEmbedding(model=embedding_model)
        ],
        vector_store=vector_store,
        docstore=SimpleDocumentStore(),
        docstore_strategy="upserts"
    )

    logger.info("Running ingestion pipeline.")
    #logger.info(f"Documents to reindex: {[doc.extra_info for doc in documents_to_reindex]}")
    pipeline.run(documents=documents_to_reindex)
    logger.info("Indexing complete.")
